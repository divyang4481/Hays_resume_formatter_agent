import sys
sys.path.insert(0, '.')
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

docx_path = Path('formatted_resume_output.docx')
doc = Document(str(docx_path))
NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

body_el = doc._element.find(f'.//{{{NS}}}body')

print(f"File: {docx_path} ({docx_path.stat().st_size} bytes)")
print(f"doc.paragraphs count: {len(doc.paragraphs)}")
print()

# Scan direct body children showing ALL text including instrText
print("=== DIRECT BODY CHILDREN ===")
para_idx = 0
for child in body_el:
    tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
    if tag == 'p':
        # get all w:t text
        t_text = ''.join(t.text or '' for t in child.findall(f'.//{{{NS}}}t')).strip()
        # get instrText
        instr = ''.join(t.text or '' for t in child.findall(f'.//{{{NS}}}instrText')).strip()
        has_fc = bool(child.findall(f'.//{{{NS}}}fldChar'))
        runs = child.findall(f'.//{{{NS}}}r')

        if t_text or instr:
            display = t_text[:70] if t_text else f'[INSTR: {instr[:50]}]'
            print(f"  p[{para_idx:03d}] runs={len(runs)} fldChar={has_fc} | {display}")
        else:
            print(f"  p[{para_idx:03d}] runs={len(runs)} fldChar={has_fc} | [EMPTY]")
        para_idx += 1
    elif tag == 'tbl':
        rows = child.findall(f'.//{{{NS}}}tr')
        print(f"  [TBL] rows={len(rows)}")
    elif tag == 'sectPr':
        print(f"  [sectPr]")
