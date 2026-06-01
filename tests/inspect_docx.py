"""Quick inspection of the rendered DOCX to verify work_experience and education were injected."""
import sys
from io import BytesIO
from pathlib import Path

sys.path.insert(0, '.')
from docx import Document

docx_path = Path('formatted_resume_v2.docx')
if not docx_path.exists():
    print('File not found:', docx_path)
    sys.exit(1)

doc = Document(str(docx_path))
print(f'File size: {docx_path.stat().st_size} bytes')
print(f'Total paragraphs: {len(doc.paragraphs)}')
print()
print('=== ALL PARAGRAPH TEXT ===')
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip()
    if txt:
        print(f'  [{i:03d}] {txt[:120]}')
