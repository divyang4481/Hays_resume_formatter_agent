from pathlib import Path

from docx import Document

from src.worker.agents.template_analysis.extractors.openxml_visual_extractor import extract_openxml_visual_evidence
from src.worker.agents.template_analysis.xml_parser import extract_fields_from_docx


def test_extract_fields_from_docx_detects_chevron_placeholders(tmp_path: Path):
    doc = Document()
    doc.add_paragraph("Candidate Name <<Candidate Name>>")
    doc.add_paragraph("Current Role <<Current Position>>")
    file_path = tmp_path / "chevron_template.docx"
    doc.save(file_path)

    fields = extract_fields_from_docx(file_path)
    names = {field["name"] for field in fields}
    tokens = {field["token"] for field in fields}

    assert "candidate_name" in names
    assert "current_position" in names
    assert "<<Candidate Name>>" in tokens
    assert "<<Current Position>>" in tokens


def test_openxml_visual_evidence_detects_chevron_placeholders():
    doc = Document()
    doc.add_paragraph("Candidate Name <<Candidate Name>>")
    doc.add_paragraph("Current Role <<Current Position>>")

    from io import BytesIO

    buffer = BytesIO()
    doc.save(buffer)
    visual_model = extract_openxml_visual_evidence(buffer.getvalue())

    public_tokens = [token.public_token for block in visual_model.blocks for token in block.tokens]

    assert "<<Candidate Name>>" in public_tokens
    assert "<<Current Position>>" in public_tokens
