from __future__ import annotations

from io import BytesIO
from pypdf import PdfReader
from docx import Document


def extract_text_from_bytes(content: bytes, filename: str | None = None) -> str:
    """
    Extracts plain text from document bytes based on file format/extension.
    Supports PDF, DOCX, and Plain Text.
    """
    # Detect based on filename or signature
    ext = ""
    if filename:
        ext = filename.lower().split(".")[-1]
    
    # Check signature magic bytes if extension not matching or unknown
    if not ext:
        if content.startswith(b"%PDF"):
            ext = "pdf"
        elif content.startswith(b"PK\x03\x04"): # Zip/Docx format
            ext = "docx"
        else:
            ext = "txt"

    if ext == "pdf":
        try:
            reader = PdfReader(BytesIO(content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")

    elif ext == "docx":
        try:
            doc = Document(BytesIO(content))
            text_parts = []
            for p in doc.paragraphs:
                if p.text.strip():
                    text_parts.append(p.text)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")

    else:
        # Fallback to plain text
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            try:
                return content.decode("latin-1")
            except Exception as e:
                raise ValueError(f"Failed to decode text: {e}")
