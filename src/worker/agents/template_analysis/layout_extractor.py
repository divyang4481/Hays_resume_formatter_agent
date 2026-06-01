from __future__ import annotations

from io import BytesIO
import re
from typing import Any

from docx import Document


def _is_heading(text: str, style_name: str) -> bool:
    t = (text or "").strip()
    s = (style_name or "").lower()
    return bool(t) and ("heading" in s or t.isupper())


def _extract_placeholder(text: str) -> str | None:
    m = re.search(r"\[[^\]]+\]|<<[^>]+>>", text or "")
    return m.group(0).strip() if m else None


def extract_layout_blocks_from_docx(docx_bytes: bytes) -> dict:
    doc = Document(BytesIO(docx_bytes))
    blocks: list[dict[str, Any]] = []
    repeat_groups: list[dict[str, Any]] = []
    current_heading = ""
    order_index = 0
    occurrence = 0

    def add_block(**kwargs: Any) -> None:
        nonlocal order_index, occurrence
        order_index += 1
        occurrence += 1
        b = {
            "block_id": f"b{order_index:03d}",
            "location": "body",
            "xml_file": "word/document.xml",
            "container_type": "paragraph",
            "section_heading": current_heading,
            "label_text": "",
            "placeholder_text": None,
            "mergefield_name": None,
            "raw_token": "",
            "paragraph_text": "",
            "table_index": None,
            "row_index": None,
            "cell_index": None,
            "paragraph_index": None,
            "is_bullet": False,
            "is_heading": False,
            "occurrence_index": occurrence,
            "order_index": order_index,
            "style_name": "",
            "evidence_text": "",
        }
        b.update(kwargs)
        blocks.append(b)

    for p_idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        style = (para.style.name if para.style else "")
        if not text:
            continue
        if _is_heading(text, style):
            current_heading = text
        ph = _extract_placeholder(text)
        if ph or "MERGEFIELD" in text.upper():
            add_block(
                container_type="paragraph",
                section_heading=current_heading,
                label_text=re.sub(r"\[[^\]]+\]", "", text).strip(" :\t"),
                placeholder_text=ph,
                raw_token=ph or text,
                paragraph_text=text,
                paragraph_index=p_idx,
                is_bullet=bool(para._p.xpath('.//w:numPr')),
                is_heading=_is_heading(text, style),
                style_name=style,
                evidence_text=text,
            )

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_text = " ".join((c.text or "").strip() for c in row.cells).strip()
            if not row_text:
                continue
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = (para.text or "").strip()
                    if not text:
                        continue
                    style = (para.style.name if para.style else "")
                    if _is_heading(text, style):
                        current_heading = text
                    ph = _extract_placeholder(text)
                    if ph or "MERGEFIELD" in text.upper() or _extract_placeholder(row_text):
                        label = re.sub(r"\[[^\]]+\]", "", row_text).strip(" :\t")
                        add_block(
                            container_type="table_cell",
                            section_heading=current_heading,
                            label_text=label,
                            placeholder_text=ph or _extract_placeholder(row_text),
                            raw_token=ph or _extract_placeholder(row_text) or text,
                            paragraph_text=text,
                            table_index=t_idx,
                            row_index=r_idx,
                            cell_index=c_idx,
                            paragraph_index=p_idx,
                            is_bullet=bool(para._p.xpath('.//w:numPr')),
                            is_heading=_is_heading(text, style),
                            style_name=style,
                            evidence_text=row_text,
                        )

    return _extract_from_xml_fields(docx_bytes)


from pathlib import Path
import tempfile
from src.worker.agents.template_analysis.xml_parser import extract_fields_from_docx
from src.worker.agents.template_analysis.logical_field_grouper import normalize_mergefield_name

def _infer_region(ev: dict, fld: dict) -> tuple[str|None,str|None,str|None]:
    token = str(fld.get("token") or "")
    label = str(ev.get("label_text") or fld.get("source_hint") or "")
    heading = str(ev.get("section_heading") or "")
    lower_heading = heading.strip().lower()
    normalized_name = normalize_mergefield_name(token, label)
    if normalized_name.startswith("presenter_"):
        return ("presenter_footer", "presenter_metadata", "PRESENTER")

    if ev.get("table_index") is not None and ev.get("row_index") is not None and ev.get("cell_index") is not None and label:
        if lower_heading in {"skills", "key skills", "professional qualifications"}:
            heading = ""
        return ("label_value_table", "label_value_pair", heading or None)

    if str(ev.get("block_type") or "") == "bullet_placeholder":
        return ("bullet_list_section", "list_item_placeholder", heading or None)

    return (None, None, heading or None)

# fallback extractor for fld codes
def _extract_from_xml_fields(docx_bytes: bytes):
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        f.write(docx_bytes)
        fp=Path(f.name)
    try:
        fields=extract_fields_from_docx(fp)
    finally:
        fp.unlink(missing_ok=True)
    blocks=[]
    for i, fld in enumerate(fields, start=1):
        ev=fld.get("template_evidence") or {}
        region_type, block_role, section_heading = _infer_region(ev, fld)
        blocks.append({
            "block_id": ev.get("block_id", f"b{i:03d}"),
            "location": "body",
            "xml_file": "word/document.xml",
            "container_type": "paragraph",
            "section_heading": section_heading or "",
            "label_text": ev.get("label_text", fld.get("source_hint","")),
            "placeholder_text": ev.get("placeholder_text") or fld.get("token"),
            "mergefield_name": None,
            "raw_token": fld.get("token"),
            "paragraph_text": ev.get("label_text", ""),
            "value_text": ev.get("cell_text") or fld.get("token"),
            "row_text": ev.get("row_text", ""),
            "cell_text": ev.get("cell_text", ""),
            "left_cell_text": ev.get("left_cell_text"),
            "right_cell_text": ev.get("right_cell_text"),
            "table_index": ev.get("table_index"),"row_index": ev.get("row_index"),"cell_index": ev.get("cell_index"),"paragraph_index": None,
            "is_bullet": bool((fld.get("context") or {}).get("is_bullet")),
            "is_heading": False,
            "occurrence_index": ev.get("occurrence_index", i),
            "order_index": i,
            "style_name": ((fld.get("context") or {}).get("style") or ""),
            "region_type": region_type,
            "block_role": block_role,
            "layout_context": {"heading": ev.get("section_heading", ""), "block_type": ev.get("block_type", "")},
            "evidence_text": ev.get("label_text") or fld.get("source_hint", ""),
        })
    return {"blocks": blocks, "repeat_groups": []}
