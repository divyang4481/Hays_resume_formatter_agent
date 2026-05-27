from __future__ import annotations

from io import BytesIO
import re
from pathlib import Path
from datetime import datetime, timezone
from typing import Any, TypedDict
from uuid import uuid4
import zipfile

from docx import Document
from langgraph.graph import END, StateGraph

from src.shared.models import GraphResult, JobStatus
from src.shared.storage import object_store
from src.worker.agentic_core import AgenticCore


class TemplateAnalysisState(TypedDict):
    template_id: str
    template_name: str
    template_object_key: str
    fields: list[dict[str, Any]]


def _normalize_field_name(token: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9]+", "_", token.strip()).strip("_").lower()
    return normalized or "unknown_field"


def _infer_field_type_from_token(name: str, template_token: str) -> tuple[str, str, str, bool]:
    token = template_token.lower()
    if re.search(r"\[\d+\]|\[\]|\{\{#|/\}\}", token):
        return "array_object", "extract repeated structured entries for this token", "section_block", False
    if "." in token:
        return "array_object", "extract nested object value for this token", "plain_text", False
    if token.count("_") >= 2:
        return "array", "extract list-like values for this token", "comma_separated", False
    return "scalar", "extract value matching this template token", "plain_text", False


# ---------------------------------------------------------------------------
# XML-level injection details scanner
# ---------------------------------------------------------------------------

def _build_injection_details_map(docx_bytes: bytes) -> dict[str, dict[str, Any]]:
    """Scan the DOCX XML to build an injection_details map per normalized field name.

    Returns:
        { normalized_field_name: injection_details_dict }

    injection_details_dict keys:
        injection_type: "mergefield" | "text_placeholder" | "table_merge_row" | "handlebars"
        mergefield_name: str  (only for mergefield)
        placeholder_text: str  (for text_placeholder)
        table_name: str  (for table_merge_row)
        sub_field_tokens: list[str]  (for table_merge_row, the sub-field MERGEFIELD names)
        locations: list[dict]  (where in the doc: body/header/footer, table/paragraph context)
    """
    details: dict[str, dict[str, Any]] = {}

    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        # Scan all word/*.xml files for tokens
        xml_files = [n for n in archive.namelist() if n.startswith("word/") and n.endswith(".xml")]
        for xml_file in xml_files:
            xml_content = archive.read(xml_file).decode("utf-8", errors="ignore")
            location_ctx = "header" if "header" in xml_file else ("footer" if "footer" in xml_file else "body")

            # --- MERGEFIELD tokens ---
            for match in re.finditer(r"MERGEFIELD\s+([a-zA-Z0-9_:]+)", xml_content, re.IGNORECASE):
                mf_name = match.group(1).strip()
                norm = _normalize_field_name(mf_name)
                # Capture a short XML preview around the match for context
                start = match.start()
                preview = xml_content[max(0, start-80): start+200].replace('\n', ' ')
                if norm not in details:
                    details[norm] = {
                        "injection_type": "mergefield",
                        "mergefield_name": mf_name,
                        "locations": [],
                    }
                details[norm]["locations"].append({"context": location_ctx, "xml_file": xml_file, "preview": preview})

            # --- TableStart / TableEnd tokens ---
            for match in re.finditer(r"TableStart:([a-zA-Z0-9_]+)", xml_content, re.IGNORECASE):
                tbl_name = match.group(1).strip()
                norm = _normalize_field_name(f"table_start_{tbl_name}")
                if norm not in details:
                    details[norm] = {
                        "injection_type": "table_merge_row",
                        "table_name": tbl_name,
                        "sub_field_tokens": [],
                        "locations": [],
                    }
                # Capture preview snippet around the TableStart token
                start = match.start()
                preview = xml_content[max(0, start-80): start+200].replace('\n', ' ')
                details[norm]["locations"].append({"context": location_ctx, "xml_file": xml_file, "preview": preview})
                # Collect sub-field MERGEFIELD names that appear near this TableStart
                # Look in same xml_file for nearby MERGEFIELDs (within 3000 chars)
                start_pos = match.start()
                nearby = xml_content[start_pos: start_pos + 3000]
                sub_mf = re.findall(r"MERGEFIELD\s+([a-zA-Z0-9_:]+)", nearby, re.IGNORECASE)
                for sf in sub_mf:
                    if sf not in details[norm]["sub_field_tokens"]:
                        details[norm]["sub_field_tokens"].append(sf)

            # --- Handlebars {{field_name}} tokens ---
            for match in re.finditer(r"\{\{\s*([a-zA-Z0-9_. -]+)\s*\}\}", xml_content):
                hb_name = match.group(1).strip()
                norm = _normalize_field_name(hb_name)
                if norm not in details:
                    details[norm] = {
                        "injection_type": "handlebars",
                        "placeholder_text": f"{{{{{hb_name}}}}}",
                        "locations": [],
                    }
                # Capture a preview snippet around the handlebars token
                start = match.start()
                preview = xml_content[max(0, start-80): start+200].replace('\n', ' ')
                details[norm]["locations"].append({"context": location_ctx, "xml_file": xml_file, "preview": preview})

    # --- Bracket [placeholder] tokens — scan paragraph plain text ---
    try:
        doc = Document(BytesIO(docx_bytes))

        def _scan_para_text(para: Any, ctx: str) -> None:
            for match in re.finditer(r"\[\s*([^\]\r\n]{2,100})\s*\]", para.text):
                raw = match.group(1).strip()
                if any(c in raw for c in ["<", ">", "=", "/"]):
                    continue
                norm = _normalize_field_name(raw)
                if norm not in details:
                    details[norm] = {
                        "injection_type": "text_placeholder",
                        "placeholder_text": f"[{raw}]",
                        "locations": [],
                    }
                # Capture a short preview of the surrounding paragraph text
                preview = para.text[:200].replace('\n', ' ')
                loc = {"context": ctx, "paragraph_preview": preview}
                details[norm]["locations"].append(loc)

        for para in doc.paragraphs:
            _scan_para_text(para, "body")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _scan_para_text(para, "body_table")
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    _scan_para_text(para, "header")
                for t in section.header.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _scan_para_text(para, "header_table")
            if section.footer:
                for para in section.footer.paragraphs:
                    _scan_para_text(para, "footer")
                for t in section.footer.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                _scan_para_text(para, "footer_table")
    except Exception as e:
        print(f"[InjectionScan] Warning scanning bracket placeholders: {e}")

    return details


def _extract_template_tokens(docx_bytes: bytes) -> list[tuple[str, str]]:
    tokens: list[tuple[str, str]] = []
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        xml_files = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
        for xml_file in xml_files:
            xml_content = archive.read(xml_file).decode("utf-8", errors="ignore")

            for match in re.findall(r"\{\{\s*([a-zA-Z0-9_. -]+)\s*\}\}", xml_content):
                tokens.append((match.strip(), f"{{{{{match.strip()}}}}}"))

            for match in re.findall(r"MERGEFIELD\s+([a-zA-Z0-9_:]+)", xml_content, flags=re.IGNORECASE):
                tokens.append((match.strip(), f"MERGEFIELD {match.strip()}"))

            for match in re.findall(r"(?:<<|«|&lt;&lt;)\s*([a-zA-Z0-9_.:-]+)\s*(?:>>|»|&gt;&gt;)", xml_content):
                tokens.append((match.strip(), f"<<{match.strip()}>>"))

            for match in re.findall(r"w:alias w:val=\"([^\"]+)\"", xml_content):
                tokens.append((match.strip(), match.strip()))

            for match in re.findall(r"TableStart:([a-zA-Z0-9_:]+)", xml_content, flags=re.IGNORECASE):
                tokens.append((match.strip(), f"TableStart:{match.strip()}"))

            for match in re.findall(r"\[\s*([^\]\r\n]{2,100})\s*\]", xml_content):
                raw_match = match.strip()
                if not any(x in raw_match for x in ["<", ">", "=", "/"]):
                    tokens.append((raw_match, f"[{raw_match}]"))

    # Extract plain text bracketed placeholders
    try:
        doc = Document(BytesIO(docx_bytes))

        def extract_from_text(text: str):
            for match in re.findall(r"\[\s*([^\]\r\n]{2,100})\s*\]", text):
                raw_match = match.strip()
                tokens.append((raw_match, f"[{raw_match}]"))

        for paragraph in doc.paragraphs:
            extract_from_text(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        extract_from_text(paragraph.text)

        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    extract_from_text(p.text)
                for t in section.header.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                extract_from_text(p.text)
            if section.footer:
                for p in section.footer.paragraphs:
                    extract_from_text(p.text)
                for t in section.footer.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                extract_from_text(p.text)
    except Exception as e:
        print(f"Error scanning document text: {e}")

    counts: dict[str, int] = {}
    output: list[tuple[str, str]] = []
    for name, original in tokens:
        key = _normalize_field_name(name)
        counts[key] = counts.get(key, 0) + 1
        out_name = key if counts[key] == 1 else f"{key}_{counts[key]}"
        output.append((out_name, original))

    return output


def _extract_template_text(docx_bytes: bytes) -> str:
    doc = Document(BytesIO(docx_bytes))
    lines: list[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)

    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                if p.text.strip():
                    lines.append(p.text.strip())
            for t in section.header.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text.strip())
        if section.footer:
            for p in section.footer.paragraphs:
                if p.text.strip():
                    lines.append(p.text.strip())
            for t in section.footer.tables:
                for row in t.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            lines.append(cell.text.strip())

    return "\n".join(lines)


def _default_fields() -> list[dict[str, Any]]:
    return []


def _field_has_evidence(field: dict[str, Any], template_text_lower: str, token_values: set[str]) -> bool:
    token = str(field.get("template_token") or "").strip()
    if token and token in token_values:
        return True
    if token and token.lower() in template_text_lower:
        return True
    if token and "macrobutton" in token.lower():
        macro_match = re.search(r"\[([^\]]+)\]", token)
        if macro_match and f"[{macro_match.group(1).strip()}]".lower() in template_text_lower:
            return True
    inj = field.get("injection_details") or {}
    if isinstance(inj, dict):
        placeholder = str(inj.get("placeholder_text") or "").strip()
        if placeholder and (placeholder in token_values or placeholder.lower() in template_text_lower):
            return True
        mergefield_name = str(inj.get("mergefield_name") or "").strip()
        if mergefield_name and (
            f"mergefield {mergefield_name}".lower() in template_text_lower
            or f"<<{mergefield_name}>>".lower() in template_text_lower
        ):
            return True
    return False


def _xml_field_to_manifest_field(xml_field: dict[str, Any], injection_map: dict[str, dict[str, Any]]) -> dict[str, Any]:
    field_name = _normalize_field_name(str(xml_field.get("name") or ""))
    token = str(xml_field.get("token") or f"[{field_name}]")
    field_type = "array" if str(xml_field.get("type") or "scalar") == "array" else "scalar"
    hint = str(xml_field.get("source_hint") or field_name.replace("_", " "))
    inj = injection_map.get(field_name)
    if not inj:
        inj = {
            "injection_type": "mergefield" if token.upper().startswith("MERGEFIELD ") else "text_placeholder",
            "mergefield_name": token[len("MERGEFIELD "):].strip() if token.upper().startswith("MERGEFIELD ") else None,
            "placeholder_text": token if token.startswith("[") else None,
            "locations": [],
        }
    return {
        "name": field_name,
        "field_type": field_type,
        "source_classification": "resume_fact",
        "source_hint": hint,
        "template_token": token,
        "required": True,
        "formatting_hint": "bullet_list" if field_type == "array" else "plain_text",
        "injection_details": inj,
        "semantic_contract": {"business_meaning": "", "resume_search_intent": "", "acceptable_sources": [], "do_not_infer": []},
        "extraction_contract": {"llm_output_key": field_name, "value_shape": field_type, "evidence_required": True, "mapping_hint": hint},
        "render_contract": {"render_strategy": "mergefield_replace" if token.upper().startswith("MERGEFIELD ") else "placeholder_replace", "anchor_token": token, "formatting": {}, "empty_value_policy": "remove_placeholder"},
        "validation_contract": {"required": True, "min_confidence": 0.65, "missing_policy": "mark_missing_do_not_generate_fake_data"},
        "source_block_ids": xml_field.get("source_block_ids") or [f"xml_{field_name}"],
        "template_evidence": xml_field.get("template_evidence") or {"template_token": token, "source_hint": hint},
    }


def _infer_fields(state: TemplateAnalysisState) -> TemplateAnalysisState:
    try:
        # Load the DOCX bytes from the object store
        template_bytes = object_store.get_bytes(state["template_object_key"])
        # 1️⃣ Extract tokens via existing regex-based scanner (used for LLM inference)
        extracted = _extract_template_tokens(template_bytes)
        template_text = _extract_template_text(template_bytes)
        # 2️⃣ Build injection details map (existing logic)
        injection_map = _build_injection_details_map(template_bytes)
        print(f"[InjectionScan] Found {len(injection_map)} injectable tokens: {list(injection_map.keys())}")
        # 3️⃣ ALSO run the new XML parser to get concrete field definitions
        from src.worker.agents.template_analysis.xml_parser import extract_fields_from_docx
        import tempfile
        # Write transient DOCX bytes to a system temp file for parser path compatibility
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
            tmp_file.write(template_bytes)
            tmp_path = Path(tmp_file.name)
        try:
            xml_fields = extract_fields_from_docx(tmp_path)
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except Exception:
                pass
        # Merge XML‑derived fields into `extracted` list, avoiding duplicates
        # Convert to dict keyed by normalized name for easy merge
        norm_to_token = {name: token for name, token in extracted}
        for field in xml_fields:
            norm_name = field["name"]
            # If the field is already present from token extraction, keep the richer LLM token
            if norm_name not in norm_to_token:
                norm_to_token[norm_name] = field["token"]
        # Re‑assemble the extracted list preserving order (XML fields first)
        extracted = [(field["name"], field["token"]) for field in xml_fields] + list(norm_to_token.items())
    except Exception as e:
        print(f"[InjectionScan] Error: {e}")
        extracted = []
        template_text = ""
        injection_map = {}

    llm_tokens = [{"name": name, "template_token": token} for name, token in extracted]

    agentic_core = AgenticCore()
    try:
        inferred = agentic_core.infer_template_manifest_fields(
            template_name=state["template_name"],
            tokens=llm_tokens,
            template_text=template_text,
            use_strong_model=False,
        )
    except Exception:
        inferred = []

    if inferred:
        normalized_fields: list[dict[str, Any]] = []
        for field in inferred:
            raw_name = str(field.get("name", ""))
            if not raw_name:
                continue
            field_name = _normalize_field_name(raw_name)
            fallback_token = str(field.get("template_token") or f"{{{{{field_name}}}}}")
            fallback_type, fallback_hint, fallback_format, fallback_required = _infer_field_type_from_token(
                field_name, fallback_token
            )
            entry: dict[str, Any] = {
                "name": field_name,
                "field_type": str(field.get("field_type") or fallback_type),
                "source_classification": str(field.get("source_classification") or "resume_fact"),
                "source_hint": str(field.get("source_hint") or fallback_hint),
                "template_token": fallback_token,
                "required": bool(field.get("required", fallback_required)),
                "formatting_hint": str(field.get("formatting_hint") or fallback_format),
            }

            # Attach injection_details from the XML scan
            # Try direct match first, then try matching by mergefield name in the token
            inj = injection_map.get(field_name)
            if inj is None:
                # Try to derive from template_token
                token_upper = fallback_token.upper()
                if token_upper.startswith("MERGEFIELD "):
                    mf_name = fallback_token[len("MERGEFIELD "):].strip()
                    mf_norm = _normalize_field_name(mf_name)
                    inj = injection_map.get(mf_norm)
                    if inj is None:
                        # Build it directly from the token
                        inj = {"injection_type": "mergefield", "mergefield_name": mf_name, "locations": []}
                elif fallback_token.upper().startswith("TABLESTART:"):
                    tbl_name = fallback_token[len("TABLESTART:"):].strip()
                    inj = injection_map.get(_normalize_field_name(f"table_start_{tbl_name}"))
                    if inj is None:
                        inj = {"injection_type": "table_merge_row", "table_name": tbl_name, "sub_field_tokens": [], "locations": []}
                elif fallback_token.startswith("["):
                    inj = {"injection_type": "text_placeholder", "placeholder_text": fallback_token, "locations": []}
                elif fallback_token.startswith("{{"):
                    inj = {"injection_type": "handlebars", "placeholder_text": fallback_token, "locations": []}
                else:
                    inj = {"injection_type": "text_placeholder", "placeholder_text": fallback_token, "locations": []}

            entry["injection_details"] = inj
            entry["semantic_contract"] = field.get("semantic_contract", {"business_meaning": "", "resume_search_intent": "", "acceptable_sources": [], "do_not_infer": []})
            entry["extraction_contract"] = field.get("extraction_contract", {"llm_output_key": field_name, "value_shape": entry["field_type"], "evidence_required": True, "mapping_hint": entry["source_hint"]})
            entry["render_contract"] = field.get("render_contract", {"render_strategy": "mergefield_replace" if inj.get("injection_type") == "mergefield" else "placeholder_replace", "anchor_token": fallback_token, "formatting": {}, "empty_value_policy": "remove_placeholder"})
            entry["validation_contract"] = field.get("validation_contract", {"required": entry["required"], "min_confidence": 0.65, "missing_policy": "mark_missing_do_not_generate_fake_data"})
            entry["source_block_ids"] = field.get("source_block_ids") or [f"legacy_{field_name}"]
            entry["template_evidence"] = field.get("template_evidence") or {"template_token": fallback_token, "source_hint": entry["source_hint"]}

            # Preserve sub_fields for array_object types
            sub_fields = field.get("sub_fields")
            if sub_fields and isinstance(sub_fields, list):
                entry["sub_fields"] = sub_fields

            normalized_fields.append(entry)

        token_values = {t.get("template_token") for t in llm_tokens if t.get("template_token")}
        template_text_lower = template_text.lower()
        normalized_fields = [
            f
            for f in normalized_fields
            if f.get("name") != "candidate_own_cv"
            or any(phrase in template_text_lower for phrase in ["candidate's own cv", "candidate cv", "paste the candidate's own cv", "original cv"])
        ]
        normalized_fields = [f for f in normalized_fields if _field_has_evidence(f, template_text_lower, token_values)]

        # Ensure deterministic XML-discovered placeholders are not dropped when LLM under-captures.
        existing_names = {f.get("name") for f in normalized_fields}
        for xf in xml_fields:
            xf_name = _normalize_field_name(str(xf.get("name") or ""))
            if not xf_name or xf_name in existing_names:
                continue
            if xf_name in {"type_text", "type_text_item", "unknown_field"}:
                continue
            candidate = _xml_field_to_manifest_field(xf, injection_map)
            if _field_has_evidence(candidate, template_text_lower, token_values):
                normalized_fields.append(candidate)
                existing_names.add(xf_name)

        if normalized_fields:
            state["fields"] = normalized_fields
            return state

    # Fallback: build from extracted tokens with injection details
    fields: list[dict[str, Any]] = []
    for normalized_name, original_token in extracted:
        field_type, source_hint, formatting_hint, required = _infer_field_type_from_token(
            normalized_name, original_token
        )
        inj = injection_map.get(normalized_name, {
            "injection_type": "text_placeholder",
            "placeholder_text": original_token,
            "locations": [],
        })
        fields.append(
            {
                "name": normalized_name,
                "field_type": field_type,
                "source_classification": "resume_fact",
                "source_hint": source_hint,
                "template_token": original_token,
                "required": required,
                "formatting_hint": formatting_hint,
                "injection_details": inj,
                "semantic_contract": {"business_meaning": "", "resume_search_intent": "", "acceptable_sources": [], "do_not_infer": []},
                "extraction_contract": {"llm_output_key": normalized_name, "value_shape": field_type, "evidence_required": True, "mapping_hint": source_hint},
                "render_contract": {"render_strategy": "mergefield_replace" if inj.get("injection_type") == "mergefield" else "placeholder_replace", "anchor_token": original_token, "formatting": {}, "empty_value_policy": "remove_placeholder"},
                "validation_contract": {"required": required, "min_confidence": 0.65, "missing_policy": "mark_missing_do_not_generate_fake_data"},
                "source_block_ids": [f"fallback_{normalized_name}"],
                "template_evidence": {"template_token": original_token, "source_hint": source_hint},
            }
        )

    state["fields"] = fields or _default_fields()
    return state


def build_template_analysis_graph():
    graph = StateGraph(TemplateAnalysisState)
    graph.add_node("infer_fields", _infer_fields)
    graph.set_entry_point("infer_fields")
    graph.add_edge("infer_fields", END)
    return graph.compile()


def run_template_analysis(template_id: str, template_name: str, template_object_key: str) -> GraphResult:
    app = build_template_analysis_graph()
    result = app.invoke(
        {
            "template_id": template_id,
            "template_name": template_name,
            "template_object_key": template_object_key,
            "fields": [],
        }
    )

    manifest = {
        "manifest_id": str(uuid4()),
        "template_id": template_id,
        "version": 2,
        "manifest_schema": "template_manifest_v2",
        "created_at": datetime.now(timezone.utc).isoformat(),
        "fields": result["fields"],
    }
    return GraphResult(status=JobStatus.COMPLETED, data=manifest)
