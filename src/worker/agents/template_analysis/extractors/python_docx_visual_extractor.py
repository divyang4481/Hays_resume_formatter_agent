from io import BytesIO
from docx import Document
from src.worker.agents.template_analysis.visual_layout_model import VisualModel, VisualTable, VisualRow, VisualCell, VisualBlock

def extract_python_docx_visual_evidence(docx_bytes: bytes) -> VisualModel:
    doc = Document(BytesIO(docx_bytes))
    model = VisualModel()

    # We mainly use python-docx to get text and basic structures, we won't reinvent token extraction
    # Since OpenXML handles tokens, this is more for structural verification

    for t_idx, table in enumerate(doc.tables):
        v_table = VisualTable(table_id=f"pd_tbl_{t_idx:03d}", table_index=t_idx)
        for r_idx, row in enumerate(table.rows):
            v_row = VisualRow(row_id=f"{v_table.table_id}_r_{r_idx:03d}", table_id=v_table.table_id, row_index=r_idx)
            row_texts = []
            for c_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                row_texts.append(cell_text)
                v_cell = VisualCell(
                    cell_id=f"{v_row.row_id}_c_{c_idx:03d}",
                    table_id=v_table.table_id,
                    row_index=r_idx,
                    cell_index=c_idx,
                    text=cell_text
                )
                v_row.cells.append(v_cell)
            v_row.row_text = " ".join(row_texts).strip()
            v_table.rows.append(v_row)
        model.tables.append(v_table)

    block_idx = 0
    for p_idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        style = para.style.name if para.style else ""
        if text:
            model.blocks.append(VisualBlock(
                block_id=f"pd_b_{block_idx:03d}",
                source="python_docx",
                page_index=None,
                order_index=block_idx,
                block_type="paragraph",
                text=text,
                style_name=style
            ))
            block_idx += 1

    return model
