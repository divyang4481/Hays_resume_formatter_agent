from __future__ import annotations

from io import BytesIO
import re
from typing import Any

from docx import Document

TOKEN_RE = re.compile(r"\[[^\]]+\]|\"[^\"]+\"")


def _is_heading(text: str, style_name: str) -> bool:
    return bool(text.strip()) and ("heading" in (style_name or "").lower() or text.isupper())


def _iter_tokens(text: str) -> list[str]:
    return [m.group(0).strip() for m in TOKEN_RE.finditer(text or "")]


def extract_python_docx_evidence(docx_bytes: bytes) -> dict[str, Any]:
    doc = Document(BytesIO(docx_bytes))
    blocks: list[dict[str, Any]] = []
    current_heading = ""
    c = 0

    def add_block(**kwargs: Any) -> None:
        nonlocal c
        c += 1
        block = {
            "source": "python_docx", "block_id": f"pd_b{c:03d}", "location": "body", "xml_file": "word/document.xml",
            "container_type": "paragraph", "section_heading": current_heading, "label_text": "", "placeholder_text": None,
            "mergefield_name": None, "raw_token": "", "table_index": None, "row_index": None, "cell_index": None,
            "paragraph_index": None, "run_index": None, "is_bullet": False, "style_name": "", "evidence_text": "",
        }
        block.update(kwargs)
        blocks.append(block)

    for p_idx, para in enumerate(doc.paragraphs):
        text = (para.text or "").strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if _is_heading(text, style):
            current_heading = text
        for t in _iter_tokens(text):
            add_block(raw_token=t, placeholder_text=t if t.startswith("[") else None, paragraph_index=p_idx, style_name=style,
                      is_bullet=bool(para._p.xpath('.//w:numPr')), evidence_text=text, label_text=re.sub(TOKEN_RE, "", text).strip(" :\t"))

    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            row_text = " ".join((c.text or "").strip() for c in row.cells).strip()
            for c_idx, cell in enumerate(row.cells):
                for p_idx, para in enumerate(cell.paragraphs):
                    text = (para.text or "").strip()
                    if not text:
                        continue
                    style = para.style.name if para.style else ""
                    if _is_heading(text, style):
                        current_heading = text
                    tokens = _iter_tokens(text) or _iter_tokens(row_text)
                    for tok in tokens:
                        add_block(container_type="table_cell", table_index=t_idx, row_index=r_idx, cell_index=c_idx, paragraph_index=p_idx,
                                  raw_token=tok, placeholder_text=tok if tok.startswith("[") else None,
                                  is_bullet=bool(para._p.xpath('.//w:numPr')), style_name=style, evidence_text=row_text or text,
                                  label_text=re.sub(TOKEN_RE, "", row_text or text).strip(" :\t"))

    # headers/footers
    for sec in doc.sections:
        for loc, container in (("header", sec.header), ("footer", sec.footer)):
            for p_idx, para in enumerate(container.paragraphs):
                text = (para.text or "").strip()
                if not text:
                    continue
                for t in _iter_tokens(text):
                    add_block(location=loc, xml_file=f"word/{loc}1.xml", raw_token=t, placeholder_text=t if t.startswith("[") else None,
                              paragraph_index=p_idx, evidence_text=text, label_text=re.sub(TOKEN_RE, "", text).strip(" :\t"))

    return {"source": "python_docx", "blocks": blocks, "warnings": []}
