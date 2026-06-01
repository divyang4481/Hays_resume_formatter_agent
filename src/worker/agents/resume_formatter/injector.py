from __future__ import annotations

import copy
import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from lxml import etree


def _normalize_field_name(token: str) -> str:
    normalized = re.sub(r"[^a-zA-Z0-9]+", "_", token.strip()).strip("_").lower()
    return normalized or "unknown_field"


def _serialize_value(value: Any, field_type: str = "scalar") -> str:
    """Convert any extracted field value to a safe string for DOCX injection."""
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, list):
        if len(value) == 0:
            return ""
        first = value[0]
        if isinstance(first, dict):
            parts: list[str] = []
            for item in value:
                row_parts = [str(v) for v in item.values() if v is not None and str(v).strip()]
                parts.append("  ".join(row_parts) if row_parts else "")
            return "\n".join(parts)
        else:
            return "\n".join(str(v) for v in value if v is not None)
    return str(value)


# ---------------------------------------------------------------------------
# Injection Strategy 1: Word Mail Merge MERGEFIELD (XML level)
# ---------------------------------------------------------------------------

def _inject_mergefields_in_element(para_elem: Any, field_map: dict[str, str]) -> None:
    """Helper to inject both complex (w:fldChar) and simple (w:fldSimple) merge fields in a paragraph element."""
    # 1. Complex fields (w:fldChar)
    runs = para_elem.findall(".//" + qn("w:r"))
    state: str | None = None     # None | "in_field" | "in_display"
    current_field: str | None = None
    display_runs: list[Any] = []

    for run in runs:
        fld_char = run.find(qn("w:fldChar"))
        instr_text = run.find(qn("w:instrText"))

        if fld_char is not None:
            fld_type = fld_char.get(qn("w:fldCharType"))
            if fld_type == "begin":
                state = "in_field"
                current_field = None
                display_runs = []
            elif fld_type == "separate":
                state = "in_display"
                display_runs = []
            elif fld_type == "end":
                if current_field and current_field in field_map and display_runs:
                    value = field_map[current_field]
                    t = display_runs[0].find(qn("w:t"))
                    if t is None:
                        t = etree.SubElement(display_runs[0], qn("w:t"))
                    t.text = value
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    for dr in display_runs[1:]:
                        t2 = dr.find(qn("w:t"))
                        if t2 is not None:
                            t2.text = ""
                state = None
                current_field = None
                display_runs = []

        elif instr_text is not None and state == "in_field":
            text = (instr_text.text or "").strip()
            m = re.search(r"MERGEFIELD\s+(\S+)", text, re.IGNORECASE)
            if m:
                current_field = m.group(1).strip().strip('"\'')

        elif state == "in_display":
            t = run.find(qn("w:t"))
            if t is not None:
                display_runs.append(run)

    # 2. Simple fields (w:fldSimple)
    fld_simples = para_elem.findall(".//" + qn("w:fldSimple"))
    for fld in fld_simples:
        instr = fld.get(qn("w:instr")) or ""
        m = re.search(r"MERGEFIELD\s+(\S+)", instr, re.IGNORECASE)
        if m:
            mf_name = m.group(1).strip().strip('"\'')
            if mf_name in field_map:
                value = field_map[mf_name]
                texts = fld.findall(".//" + qn("w:t"))
                if texts:
                    texts[0].text = value
                    texts[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                    for t in texts[1:]:
                        t.text = ""
                else:
                    r = etree.SubElement(fld, qn("w:r"))
                    t = etree.SubElement(r, qn("w:t"))
                    t.text = value
                    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")


def _row_contains_tablestart(tr_elem: Any, table_name: str) -> bool:
    for instr in tr_elem.findall(".//" + qn("w:instrText")):
        if f"TableStart:{table_name}" in (instr.text or ""):
            return True
    for fld in tr_elem.findall(".//" + qn("w:fldSimple")):
        instr = fld.get(qn("w:instr")) or ""
        if f"TableStart:{table_name}" in instr:
            return True
    return False


def _inject_table_merge_rows(doc: Document, table_data_map: dict[str, list[dict[str, str]]]) -> None:
    """Scan all tables in the document (including nested ones) and expand rows matching TableStart:<table_name> markers."""
    if not table_data_map:
        return

    # Find all w:tr elements anywhere in the document XML (handles nested tables perfectly!)
    tr_elements = doc._element.findall(".//" + qn("w:tr"))
    for tr_elem in tr_elements:
        if tr_elem.getparent() is None:
            continue
            
        matched_table_name = None
        for tname in table_data_map.keys():
            if _row_contains_tablestart(tr_elem, tname):
                matched_table_name = tname
                break
        
        if matched_table_name:
            items = table_data_map[matched_table_name]
            tr_parent = tr_elem.getparent()
            tr_index = tr_parent.index(tr_elem)
            
            # Clone and inject for each item in the list
            for item in items:
                cloned_tr = copy.deepcopy(tr_elem)
                
                local_field_map = {
                    f"TableStart:{matched_table_name}": "",
                    f"TableEnd:{matched_table_name}": "",
                    **item
                }
                
                # Also map lowercase/stripped keys for reliability
                for k, v in list(local_field_map.items()):
                    local_field_map[k.lower()] = v
                    local_field_map[k.strip()] = v
                
                for para_p in cloned_tr.findall(".//" + qn("w:p")):
                    _inject_mergefields_in_element(para_p, local_field_map)
                
                # Insert the cloned row
                tr_parent.insert(tr_index, cloned_tr)
                tr_index += 1
            
            # Remove the original template row
            tr_parent.remove(tr_elem)


def _inject_mergefields(doc: Document, field_map: dict[str, str]) -> None:
    """Replace Word MERGEFIELD values (both simple and complex) at the XML element level."""
    if not field_map:
        return

    # Process all w:p elements anywhere in the document XML (handles nested tables, content controls, and text boxes!)
    for para_el in doc._element.findall(".//" + qn("w:p")):
        _inject_mergefields_in_element(para_el, field_map)

    # Process all headers and footers XML
    for section in doc.sections:
        if section.header:
            for para_el in section.header._element.findall(".//" + qn("w:p")):
                _inject_mergefields_in_element(para_el, field_map)
        if section.footer:
            for para_el in section.footer._element.findall(".//" + qn("w:p")):
                _inject_mergefields_in_element(para_el, field_map)

    print(f"[MergeField] Applied XML-level replacement for {len(field_map)} fields: {list(field_map.keys())}")


# ---------------------------------------------------------------------------
# Injection Strategy 2: Text placeholder replacement (Macrobutton aware)
# ---------------------------------------------------------------------------

def _replace_in_paragraph(paragraph: Any, key: str, value: str) -> bool:
    """Replace a text placeholder in a paragraph. Returns True if replaced."""
    if key not in paragraph.text:
        return False

    # Fast path: placeholder fully in one run
    for run in paragraph.runs:
        if key in run.text:
            run.text = run.text.replace(key, value)
            return True

    # Slow path: placeholder split across runs — merge then redistribute
    if paragraph.runs:
        full_text = "".join(run.text for run in paragraph.runs)
        if key in full_text:
            new_text = full_text.replace(key, value)
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
            return True
    return False


def _paragraph_contains_token(para: Any, token: str) -> bool:
    """Check if a paragraph contains a token, either in plain text or inside a macrobutton instruction."""
    if token in para.text:
        return True
    p_elem = para._element
    for instr in p_elem.findall(".//" + qn("w:instrText")):
        if token in (instr.text or ""):
            return True
    return False


def _replace_token_in_paragraph(para: Any, token: str, value: str) -> bool:
    """Replace a token in a paragraph, handling both plain text and macrobutton XML fields cleanly."""
    if token in para.text:
        return _replace_in_paragraph(para, token, value)

    # Check macrobutton field replacement
    has_macrobutton = False
    p_elem = para._element
    for instr in p_elem.findall(".//" + qn("w:instrText")):
        if token in (instr.text or ""):
            has_macrobutton = True
            break
            
    if has_macrobutton:
        # Remove ALL direct children of w:p except w:pPr (paragraph properties).
        # Using findall(".//w:r") + getparent().remove() is unreliable because
        # lxml's findall returns a snapshot but remove() mutates the tree mid-iteration.
        # Clearing direct children preserves formatting (pPr) while wiping the entire
        # Macrobutton complex-field structure (fldChar begin/instrText/separate/end runs).
        pPr = p_elem.find(qn("w:pPr"))
        for child in list(p_elem):          # list() snapshots children before mutation
            p_elem.remove(child)
        if pPr is not None:
            p_elem.insert(0, pPr)           # restore paragraph properties

        # Add a new clean text run containing the replacement value
        para.add_run(value)
        return True

    return False


def _inject_text_placeholders(doc: Document, text_map: dict[str, str]) -> None:
    """Replace bracket-style and handlebars-style placeholders in paragraph text."""
    if not text_map:
        return

    # Process all w:p elements anywhere in the document XML (handles nested tables, content controls, and text boxes!)
    for para_el in doc._element.findall(".//" + qn("w:p")):
        para_obj = Paragraph(para_el, doc)
        for token, val in text_map.items():
            _replace_token_in_paragraph(para_obj, token, val)

    # Process all headers and footers XML
    for section in doc.sections:
        if section.header:
            for para_el in section.header._element.findall(".//" + qn("w:p")):
                para_obj = Paragraph(para_el, section.header)
                for token, val in text_map.items():
                    _replace_token_in_paragraph(para_obj, token, val)
        if section.footer:
            for para_el in section.footer._element.findall(".//" + qn("w:p")):
                para_obj = Paragraph(para_el, section.footer)
                for token, val in text_map.items():
                    _replace_token_in_paragraph(para_obj, token, val)

    print(f"[TextReplace] Applied text placeholder replacement for {len(text_map)} tokens")


def _is_heading(para: Any) -> bool:
    """Helper to detect if a paragraph is a visual/semantic heading in the template."""
    style_name = (para.style.name or "").lower()
    if "heading" in style_name:
        return True
    # Heuristic: short, all-caps or contains bold text
    text = para.text.strip()
    if 2 < len(text) < 50:
        if text.isupper():
            return True
        if any(run.bold for run in para.runs if run.text.strip()):
            return True
    return False


def _match_section(current_section: str, field_section: str, field_name: str) -> float:
    """Calculate a semantic score between the active template section heading and manifest field metadata."""
    cs = re.sub(r"[^a-z]+", "", current_section.lower())
    fs = re.sub(r"[^a-z]+", "", field_section.lower())
    fn = re.sub(r"[^a-z]+", "", field_name.lower())
    
    if fs and (fs in cs or cs in fs):
        return 1.0
    if fn and (fn in cs or cs in fn):
        return 0.9
        
    keywords_map = {
        "skill": ["skill", "expertise", "competenc"],
        "qualification": ["qualification", "certific", "credential", "education"],
        "experience": ["experience", "employment", "history", "career", "work"],
    }
    
    for kw, synonyms in keywords_map.items():
        cs_match = any(syn in cs for syn in synonyms)
        fs_match = any(syn in fs for syn in synonyms) or any(syn in fn for syn in synonyms)
        if cs_match and fs_match:
            return 0.8
            
    return 0.0


def _inject_array_paragraphs(
    doc: Document,
    extracted: dict[str, Any],
    manifest_fields: list[dict[str, Any]]
) -> None:
    """Replicate bullet paragraph elements N times for arrays, mapping them dynamically based on sections."""
    array_fields: list[dict] = []
    for field in manifest_fields:
        fname = field.get("name", "")
        ftype = field.get("field_type", "scalar")
        inj = field.get("injection_details") or {}
        injection_type = inj.get("injection_type", "text_placeholder")
        render_strategy = (field.get("render_contract") or {}).get("render_strategy", "")

        # Fields with explicit block mappings are handled in targeted injection.
        if field.get("source_block_ids"):
            continue

        # Table-region arrays must be rendered only by table row expansion.
        # If we expand them as generic paragraphs, long values end up in the wrong/narrow cell.
        if render_strategy == "mailmerge_table_region":
            continue
        
        if ftype in ("array", "array_object") and injection_type in ("text_placeholder", "handlebars"):
            val = extracted.get(fname)
            if not val:
                val = []
            elif not isinstance(val, list):
                val = [val]
                
            serialized_items = []
            for item in val:
                if isinstance(item, dict):
                    parts = [str(v) for v in item.values() if v is not None and str(v).strip()]
                    serialized_items.append("  ".join(parts) if parts else "")
                else:
                    serialized_items.append(str(item))
                    
            serialized_items = [s for s in serialized_items if s.strip()]
            token = inj.get("placeholder_text") or field.get("template_token") or f"{{{{{fname}}}}}"
            
            array_fields.append({
                "name": fname,
                "token": token,
                "section": field.get("layout_details", {}).get("section") or "",
                "items": serialized_items,
            })

    if not array_fields:
        return

    expanded_fields: set[str] = set()

    # Collect all w:p elements from body, header, footer in document order
    paras_to_process: list[tuple[Any, Any]] = []
    for p_el in doc._element.findall(".//" + qn("w:p")):
        paras_to_process.append((p_el, doc))
    for section in doc.sections:
        if section.header:
            for p_el in section.header._element.findall(".//" + qn("w:p")):
                paras_to_process.append((p_el, section.header))
        if section.footer:
            for p_el in section.footer._element.findall(".//" + qn("w:p")):
                paras_to_process.append((p_el, section.footer))

    current_section = ""
    i = 0
    while i < len(paras_to_process):
        para_el, parent_obj = paras_to_process[i]
        if para_el.getparent() is None:
            i += 1
            continue
            
        para = Paragraph(para_el, parent_obj)
        if _is_heading(para):
            current_section = para.text.strip()
            
        matched_field = None
        matched_token = None
        best_score = -1.0
        
        for af in array_fields:
            token = af["token"]
            tokens_to_check = [token, f"{{{{{af['name']}}}}}", f"{{{{ {af['name']} }}}}"]
            
            found_token = None
            for t in tokens_to_check:
                if t and _paragraph_contains_token(para, t):
                    found_token = t
                    break
                    
            if found_token:
                score = _match_section(current_section, af["section"], af["name"])
                if score > best_score:
                    best_score = score
                    matched_field = af
                    matched_token = found_token

        if matched_field and matched_token:
            fname = matched_field["name"]
            items = matched_field["items"]
            p_elem = para._element
            p_parent = p_elem.getparent()
            p_index = p_parent.index(p_elem)
            
            if fname not in expanded_fields:
                expanded_fields.add(fname)
                for item in items:
                    cloned_p = copy.deepcopy(p_elem)
                    cloned_para_obj = Paragraph(cloned_p, para._parent)
                    _replace_token_in_paragraph(cloned_para_obj, matched_token, item)
                    p_parent.insert(p_index, cloned_p)
                    p_index += 1
                
                p_parent.remove(p_elem)
            else:
                p_parent.remove(p_elem)
        i += 1


# ---------------------------------------------------------------------------
# Main injection dispatcher — uses injection_details from manifest
# ---------------------------------------------------------------------------

def _inject_targeted_value(doc: Document, fname: str, token: str, value: str, field: dict) -> bool:
    """Inject value selectively only in the paragraphs or table cells matching source_block_ids."""
    source_block_ids = field.get("source_block_ids") or []
    if not source_block_ids:
        return False
        
    render_strategy = field.get("render_contract", {}).get("render_strategy", "")
    is_instruction = field.get("template_evidence", {}).get("is_instruction_only", False) or \
                     "instruction" in str(field.get("template_evidence", {}).get("region_type", "")).lower() or \
                     render_strategy == "remove_instruction_text"
    region_type = str(field.get("template_evidence", {}).get("region_type", "")).lower()
    field_type = str(field.get("field_type", "scalar")).lower()
    collapse_multi_block_scalar = (
        field_type == "scalar"
        and len(source_block_ids) > 1
        and region_type == "bullet_list_section"
    )

    replaced = False
    bullet_items: list[str] = []
    if collapse_multi_block_scalar:
        # Many templates model list sections as multiple [Type text] bullet placeholders.
        # Split comma/newline-delimited text so each placeholder gets one bullet item.
        raw_parts = re.split(r"\n|\r|\u2022|\s*,\s*", value or "")
        bullet_items = [p.strip() for p in raw_parts if p and p.strip()]
        if not bullet_items and value:
            bullet_items = [value.strip()]

    def _resolve_body_paragraph_by_block_index(p_idx: int) -> Paragraph | None:
        """Resolve b_XXX index using the same top-level body paragraph semantics as visual extraction."""
        body_el = doc._element.find(".//" + qn("w:body"))
        if body_el is None:
            return None

        block_counter = 0
        for child in body_el:
            if child.tag != qn("w:p"):
                continue

            texts = [(t.text or "") for t in child.findall(".//" + qn("w:t"))]
            paragraph_text = "".join(texts).strip()
            # Include w:fldChar (complex/Macrobutton fields) in addition to simple fields
            has_tokens = bool(
                child.findall(".//" + qn("w:fldSimple")) or
                child.findall(".//" + qn("w:instrText")) or
                child.findall(".//" + qn("w:fldChar"))
            )
            if not (paragraph_text or has_tokens):
                continue

            if block_counter == p_idx:
                return Paragraph(child, doc)
            block_counter += 1

        return None

    if is_instruction:
        # Instruction/CV page block replacement: first block gets the entire value, other blocks are cleared
        first_p_replaced = False
        for block_id in source_block_ids:
            if block_id.startswith("b_"):
                parts = block_id.split("_")
                if len(parts) == 2:
                    try:
                        p_idx = int(parts[1])
                        para = _resolve_body_paragraph_by_block_index(p_idx)
                        if para is None:
                            continue

                        # Clear all existing runs/text in this block
                        for r in para._element.findall(".//" + qn("w:r")):
                            r.getparent().remove(r)

                        if not first_p_replaced:
                            para.add_run(value)
                            first_p_replaced = True
                            print(f"  [Targeted Instruction] Injected value into paragraph block {block_id}")
                        else:
                            para.text = ""
                            print(f"  [Targeted Instruction] Cleared paragraph block {block_id}")
                        replaced = True
                    except Exception as e:
                        print(f"Error replacing instruction block {block_id}: {e}")
        return replaced

    for block_pos, block_id in enumerate(source_block_ids):
        if block_id.startswith("tbl_"):
            # Table cell format: tbl_XXX_r_YYY_c_ZZZ
            parts = block_id.split("_")
            if len(parts) >= 6:
                try:
                    t_idx = int(parts[1])
                    r_idx = int(parts[3])
                    c_idx = int(parts[5])
                    
                    # Robust XML-based visual table cell locator (handles nested tables perfectly!)
                    body_el = doc._element.find(".//" + qn("w:body"))
                    if body_el is not None:
                        top_level_tables = [child for child in body_el if child.tag == qn("w:tbl")]
                        if t_idx < len(top_level_tables):
                            table_elem = top_level_tables[t_idx]
                            row_elements = table_elem.findall("./" + qn("w:tr"))
                            if r_idx < len(row_elements):
                                row_elem = row_elements[r_idx]
                                cell_elements = row_elem.findall("./" + qn("w:tc"))
                                if c_idx < len(cell_elements):
                                    cell_elem = cell_elements[c_idx]
                                    para_elements = cell_elem.findall(".//" + qn("w:p"))
                                    for para_el in para_elements:
                                        para_obj = Paragraph(para_el, doc)
                                        if _replace_token_in_paragraph(para_obj, token, value):
                                            replaced = True
                                    # Also handle MERGEFIELD-type tokens (e.g. «CandidateFullName»)
                                    # _replace_token_in_paragraph only searches visible text;
                                    # MERGEFIELDs live inside w:instrText and need this second pass.
                                    mf_map = {token: value, token.lower(): value}
                                    for para_el in para_elements:
                                        _inject_mergefields_in_element(para_el, mf_map)
                except Exception as e:
                    print(f"Error doing targeted table cell injection for {block_id}: {e}")
        elif block_id.startswith("pd_tbl_"):
            # pd_tbl format: pd_tbl_XXX
            pass
        elif block_id.startswith("b_"):
            # Paragraph format: b_XXX
            parts = block_id.split("_")
            if len(parts) == 2:
                try:
                    p_idx = int(parts[1])
                    para = _resolve_body_paragraph_by_block_index(p_idx)
                    if para is None:
                        continue

                    if collapse_multi_block_scalar:
                        replacement = ""
                        if block_pos < len(bullet_items):
                            if block_pos == len(source_block_ids) - 1 and len(bullet_items) > len(source_block_ids):
                                replacement = "\n".join(bullet_items[block_pos:])
                            else:
                                replacement = bullet_items[block_pos]

                        # Fill one bullet placeholder per block and clear unused placeholders.
                        if _replace_token_in_paragraph(para, token, replacement):
                            replaced = True
                        continue

                    replaced_here = _replace_token_in_paragraph(para, token, value)
                    if replaced_here:
                        replaced = True

                    # Also handle MERGEFIELD-type tokens on this paragraph element.
                    _inject_mergefields_in_element(
                        para._element,
                        {token: value, token.lower(): value},
                    )
                except Exception as e:
                    print(f"Error doing targeted paragraph injection for {block_id}: {e}")

    # Fallback for table-driven fields when strict tbl/row/cell addressing drifts across template variants.
    # Only apply to non-bracket tokens (e.g., MERGEFIELD-like names such as CandidateTown),
    # because bracket placeholders like [Type text] can be intentionally reused across sections.
    if (not replaced) and token and (not token.startswith("[")) and (" " not in token):
        fallback_map = {token: value, token.lower(): value}
        for para_el in doc._element.findall(".//" + qn("w:p")):
            para_obj = Paragraph(para_el, doc)
            replaced_here = _replace_token_in_paragraph(para_obj, token, value)
            _inject_mergefields_in_element(para_el, fallback_map)
            if replaced_here:
                replaced = True
        if replaced:
            print(f"  [Targeted Fallback] injected {fname} by token scan using '{token}'")

    return replaced


def _get_item_value_for_subfield(item: dict[str, Any], sub_name: str) -> Any:
    """Resolve a subfield value from an extracted row item with tolerant key matching."""
    if sub_name in item:
        return item.get(sub_name)

    wanted = _normalize_field_name(sub_name)
    for k, v in item.items():
        if _normalize_field_name(str(k)) == wanted:
            return v
    return None


def _replace_any_token_in_paragraph(para: Paragraph, tokens: list[str], value: str) -> bool:
    for t in tokens:
        if t and _replace_token_in_paragraph(para, t, value):
            return True
    return False


def _token_variants(token: str, sub_name: str) -> list[str]:
    variants = [token]
    if token.startswith("[") and token.endswith("]"):
        variants.append(token[1:-1].strip())
    variants.extend([
        f"{{{{{sub_name}}}}}",
        f"{{{{ {sub_name} }}}}",
        f"[{sub_name}]",
        f'"{token}"',
        f"'{token}'",
    ])

    deduped: list[str] = []
    seen: set[str] = set()
    for v in variants:
        if v and v not in seen:
            deduped.append(v)
            seen.add(v)
    return deduped


def _collect_indexed_body_paragraphs(doc: Document) -> list[Paragraph]:
    """Collect non-empty/tokenized body paragraphs in document order, including nested content."""
    body_el = doc._element.find(".//" + qn("w:body"))
    if body_el is None:
        return []

    paragraphs: list[Paragraph] = []
    for p_el in body_el.findall(".//" + qn("w:p")):
        if p_el.getparent() is None:
            continue

        texts = [(t.text or "") for t in p_el.findall(".//" + qn("w:t"))]
        paragraph_text = "".join(texts).strip()
        has_tokens = bool(
            p_el.findall(".//" + qn("w:fldSimple")) or
            p_el.findall(".//" + qn("w:instrText")) or
            p_el.findall(".//" + qn("w:fldChar"))
        )
        if paragraph_text or has_tokens:
            paragraphs.append(Paragraph(p_el, doc))

    return paragraphs


def _inject_repeat_block_by_section_tokens(doc: Document, field: dict[str, Any], value: list[Any]) -> bool:
    """Fallback repeat-block injection using token occurrences within the field's section."""
    if not value:
        return False

    render_contract = field.get("render_contract") or {}
    block_tokens = render_contract.get("block_tokens") or {}
    sub_fields = field.get("sub_fields") or []
    section_heading = str((field.get("template_evidence") or {}).get("section_heading") or "").strip()

    if not sub_fields:
        return False

    all_paras = [
        Paragraph(p_el, doc)
        for p_el in doc._element.findall(".//" + qn("w:p"))
        if p_el.getparent() is not None
    ]
    if not all_paras:
        return False

    scoped_paras: list[Paragraph] = []
    if section_heading:
        in_section = False
        target_norm = re.sub(r"[^a-z0-9]+", "", section_heading.lower())
        for para in all_paras:
            txt = (para.text or "").strip()
            if _is_heading(para):
                heading_norm = re.sub(r"[^a-z0-9]+", "", txt.lower())
                if target_norm and target_norm in heading_norm:
                    in_section = True
                elif in_section:
                    break

            if in_section:
                scoped_paras.append(para)

    if not scoped_paras:
        scoped_paras = all_paras

    cursors: dict[str, int] = {sf.get("name", ""): 0 for sf in sub_fields if sf.get("name")}
    injected_any = False

    for row in value:
        if not isinstance(row, dict):
            continue
        for sf in sub_fields:
            sf_name = sf.get("name", "")
            if not sf_name:
                continue

            sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
            variants = _token_variants(sf_token, sf_name)
            sf_value = _get_item_value_for_subfield(row, sf_name)
            serialized = _serialize_value(sf_value, sf.get("field_type", "scalar"))
            if not serialized:
                continue

            start_idx = cursors.get(sf_name, 0)
            for idx in range(start_idx, len(scoped_paras)):
                para = scoped_paras[idx]
                if any(_paragraph_contains_token(para, v) for v in variants):
                    if _replace_any_token_in_paragraph(para, variants, serialized):
                        injected_any = True
                    cursors[sf_name] = idx + 1
                    break

    # Remove leftover template placeholders for repeat-block tokens in this section.
    for sf in sub_fields:
        sf_name = sf.get("name", "")
        if not sf_name:
            continue
        sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
        variants = _token_variants(sf_token, sf_name)
        for para in scoped_paras:
            if any(_paragraph_contains_token(para, v) for v in variants):
                if _replace_any_token_in_paragraph(para, variants, ""):
                    injected_any = True

    return injected_any


def _inject_array_targeted(doc: Document, field: dict[str, Any], value: list[Any]) -> bool:
    """Inject flat array fields into their mapped paragraph blocks, duplicating as needed."""
    if not isinstance(value, list) or not value:
        return False

    source_block_ids = field.get("source_block_ids") or []
    if not source_block_ids:
        return False

    token = field.get("template_token") or field.get("token") or f"[{field.get('name')}]"
    injected_any = False

    # Filter out empty/invalid items
    items = [_serialize_value(item, "scalar") for item in value if item is not None]
    items = [item for item in items if item.strip()]

    if not items:
        return False

    # Resolve all body paragraphs once for fast lookup
    # NOTE: Must match the visual extraction block indexing exactly — include w:fldChar
    # (complex field / Macrobutton) paragraphs so that b_XXX indices align correctly.
    body_el = doc._element.find(".//" + qn("w:body"))
    all_body_paras: list[Paragraph] = []
    if body_el is not None:
        for child in body_el:
            if child.tag != qn("w:p"):
                continue
            texts = [(t.text or "") for t in child.findall(".//" + qn("w:t"))]
            paragraph_text = "".join(texts).strip()
            has_tokens = bool(
                child.findall(".//" + qn("w:fldSimple")) or
                child.findall(".//" + qn("w:instrText")) or
                child.findall(".//" + qn("w:fldChar"))
            )
            if not (paragraph_text or has_tokens):
                continue
            all_body_paras.append(Paragraph(child, doc))

    def get_para_by_id(block_id: str) -> Paragraph | None:
        if block_id.startswith("b_"):
            parts = block_id.split("_")
            if len(parts) == 2:
                try:
                    p_idx = int(parts[1])
                    if 0 <= p_idx < len(all_body_paras):
                        return all_body_paras[p_idx]
                except ValueError:
                    pass
        return None

    resolved_paras: list[Paragraph] = []
    for b_id in source_block_ids:
        para = get_para_by_id(b_id)
        if para is not None:
            resolved_paras.append(para)

    if not resolved_paras:
        return False

    num_items = len(items)
    n_paras = len(resolved_paras)

    # For the first n_paras - 1 paragraphs:
    for idx in range(n_paras - 1):
        para = resolved_paras[idx]
        item_val = items[idx] if idx < num_items else ""
        if _replace_token_in_paragraph(para, token, item_val):
            injected_any = True

    # For the last paragraph:
    last_para = resolved_paras[-1]
    last_idx = n_paras - 1

    if last_idx < num_items:
        # Keep a deepcopy of the original element before replacement
        p_elem = last_para._element
        p_parent = p_elem.getparent()
        orig_p_elem = copy.deepcopy(p_elem)

        # Inject into the last para itself
        if _replace_token_in_paragraph(last_para, token, items[last_idx]):
            injected_any = True

        # Inject subsequent items by duplicating the original paragraph
        insert_after_para = last_para
        for item in items[last_idx + 1:]:
            cloned_el = copy.deepcopy(orig_p_elem)
            cloned_para = Paragraph(cloned_el, last_para._parent)
            _replace_token_in_paragraph(cloned_para, token, item)

            p_elem = insert_after_para._element
            p_parent = p_elem.getparent()
            if p_parent is not None:
                p_index = p_parent.index(p_elem)
                p_parent.insert(p_index + 1, cloned_el)
                insert_after_para = cloned_para
                injected_any = True
    else:
        # Fewer items than paragraphs. Clear the last paragraph.
        if _replace_token_in_paragraph(last_para, token, ""):
            injected_any = True

    return injected_any


def _inject_repeat_block_targeted(doc: Document, field: dict[str, Any], value: Any) -> bool:
    """Inject array_object repeat-block fields into their mapped paragraph blocks, duplicating patterns as needed."""
    if not isinstance(value, list) or not value:
        return False

    render_contract = field.get("render_contract") or {}
    block_tokens = render_contract.get("block_tokens") or {}
    repeat_items = render_contract.get("repeat_items") or []
    sub_fields = field.get("sub_fields") or []
    source_block_ids = field.get("source_block_ids") or []

    if not sub_fields:
        return False

    sub_field_type: dict[str, str] = {
        sf.get("name", ""): sf.get("field_type", "scalar") for sf in sub_fields if sf.get("name")
    }

    # Resolve all body paragraphs once for fast lookup
    # NOTE: Must match the visual extraction block indexing exactly — include w:fldChar
    # (complex field / Macrobutton) paragraphs so that b_XXX indices align correctly.
    body_el = doc._element.find(".//" + qn("w:body"))
    all_body_paras: list[Paragraph] = []
    if body_el is not None:
        for child in body_el:
            if child.tag != qn("w:p"):
                continue
            texts = [(t.text or "") for t in child.findall(".//" + qn("w:t"))]
            paragraph_text = "".join(texts).strip()
            has_tokens = bool(
                child.findall(".//" + qn("w:fldSimple")) or
                child.findall(".//" + qn("w:instrText")) or
                child.findall(".//" + qn("w:fldChar"))
            )
            if not (paragraph_text or has_tokens):
                continue
            all_body_paras.append(Paragraph(child, doc))

    def get_para_by_id(block_id: str) -> Paragraph | None:
        if block_id.startswith("b_"):
            parts = block_id.split("_")
            if len(parts) == 2:
                try:
                    p_idx = int(parts[1])
                    if 0 <= p_idx < len(all_body_paras):
                        return all_body_paras[p_idx]
                except ValueError:
                    pass
        return None

    # Step 1: Build the base pattern
    base_paras: list[Paragraph] = []
    sf_to_para_idx: dict[str, int] = {}

    if repeat_items:
        block_map = repeat_items[0]
        resolved_blocks = []
        for sf in sub_fields:
            sf_name = sf.get("name", "")
            b_id = block_map.get(sf_name)
            if b_id:
                para = get_para_by_id(b_id)
                if para is not None:
                    resolved_blocks.append((sf_name, b_id, para))

        def get_p_idx(item):
            b_id = item[1]
            try:
                return int(b_id.split("_")[1])
            except (ValueError, IndexError):
                return 0

        resolved_blocks.sort(key=get_p_idx)
        base_paras = [item[2] for item in resolved_blocks]
        for idx, item in enumerate(resolved_blocks):
            sf_name = item[0]
            sf_to_para_idx[sf_name] = idx
    else:
        for b_id in source_block_ids:
            para = get_para_by_id(b_id)
            if para is not None:
                base_paras.append(para)

        for sf in sub_fields:
            sf_name = sf.get("name", "")
            if not sf_name:
                continue
            sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
            variants = _token_variants(sf_token, sf_name)

            for idx, para in enumerate(base_paras):
                if any(_paragraph_contains_token(para, v) for v in variants):
                    sf_to_para_idx[sf_name] = idx
                    break

    if not base_paras:
        # Fallback to section-based tokens if we cannot build a base pattern
        return _inject_repeat_block_by_section_tokens(doc, field, value)

    # Step 2: Keep deep copy of original elements for cloning
    orig_elements = [copy.deepcopy(p._element) for p in base_paras]

    injected_any = False
    insert_after_para = base_paras[-1]

    # Step 3: Inject and duplicate
    for row_idx, row_item in enumerate(value):
        if not isinstance(row_item, dict):
            continue

        if row_idx == 0:
            # Reuse base_paras in-place
            for sf in sub_fields:
                sf_name = sf.get("name", "")
                if sf_name not in sf_to_para_idx:
                    continue
                pat_idx = sf_to_para_idx[sf_name]
                para = base_paras[pat_idx]
                
                sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
                variants = _token_variants(sf_token, sf_name)
                sf_value = _get_item_value_for_subfield(row_item, sf_name)
                serialized = _serialize_value(sf_value, sub_field_type.get(sf_name, "scalar"))
                
                if _replace_any_token_in_paragraph(para, variants, serialized):
                    injected_any = True
            insert_after_para = base_paras[-1]

        elif repeat_items and row_idx < len(repeat_items):
            # Reuse existing dummy block in-place
            block_map = repeat_items[row_idx]
            max_p_idx = -1
            last_para_of_block = None

            for sf in sub_fields:
                sf_name = sf.get("name", "")
                b_id = block_map.get(sf_name)
                if not b_id:
                    continue
                para = get_para_by_id(b_id)
                if para is None:
                    continue

                if b_id.startswith("b_"):
                    try:
                        p_idx = int(b_id.split("_")[1])
                        if p_idx > max_p_idx:
                            max_p_idx = p_idx
                            last_para_of_block = para
                    except ValueError:
                        pass

                sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
                variants = _token_variants(sf_token, sf_name)
                sf_value = _get_item_value_for_subfield(row_item, sf_name)
                serialized = _serialize_value(sf_value, sub_field_type.get(sf_name, "scalar"))

                if _replace_any_token_in_paragraph(para, variants, serialized):
                    injected_any = True

            if last_para_of_block:
                insert_after_para = last_para_of_block

        else:
            # Duplicate the base pattern and insert
            new_clones = []
            for pat_idx, orig_el in enumerate(orig_elements):
                cloned_el = copy.deepcopy(orig_el)
                cloned_para = Paragraph(cloned_el, base_paras[pat_idx]._parent)

                # Inject values into the cloned element
                for sf in sub_fields:
                    sf_name = sf.get("name", "")
                    if sf_to_para_idx.get(sf_name) == pat_idx:
                        sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
                        variants = _token_variants(sf_token, sf_name)
                        sf_value = _get_item_value_for_subfield(row_item, sf_name)
                        serialized = _serialize_value(sf_value, sub_field_type.get(sf_name, "scalar"))
                        _replace_any_token_in_paragraph(cloned_para, variants, serialized)

                # Insert cloned_el right after insert_after_para
                p_elem = insert_after_para._element
                p_parent = p_elem.getparent()
                if p_parent is not None:
                    p_index = p_parent.index(p_elem)
                    p_parent.insert(p_index + 1, cloned_el)
                    insert_after_para = cloned_para
                    injected_any = True

    # Step 4: Clear any unused pre-existing dummy blocks
    if repeat_items and len(value) < len(repeat_items):
        for unused_idx in range(len(value), len(repeat_items)):
            block_map = repeat_items[unused_idx]
            for sf_name, b_id in block_map.items():
                para = get_para_by_id(b_id)
                if para is not None:
                    sf = next((s for s in sub_fields if s.get("name") == sf_name), None)
                    if sf:
                        sf_token = block_tokens.get(sf_name) or sf.get("template_token") or f"[{sf_name}]"
                        variants = _token_variants(sf_token, sf_name)
                        _replace_any_token_in_paragraph(para, variants, "")

    return injected_any


def inject_data_into_docx(
    docx_bytes: bytes,
    extracted: dict[str, Any],
    manifest_fields: list[dict[str, Any]],
) -> bytes:
    """Inject extracted resume values into the DOCX template using per-field injection strategy."""
    doc = Document(BytesIO(docx_bytes))

    mergefield_map: dict[str, str] = {}   # mergefield_name_in_xml → value
    text_placeholder_map: dict[str, str] = {}  # placeholder_text → value
    table_data_map: dict[str, list[dict[str, str]]] = {} # table_name -> list of dict values mapped by xml sub-field
    populated_fields: set[str] = set()

    # 1. Targeted injection pass first to handle visual block mappings (like multi-use [Type text])
    targeted_injected_fields = set()
    for field in manifest_fields:
        fname = field.get("name", "")
        ftype = field.get("field_type", "scalar")
        value = extracted.get(fname)
        
        # User constraint: if value is empty or not available, do not inject/replace!
        if value is None or value == "" or value == [] or (isinstance(value, str) and not value.strip()):
            print(f"  [Skip Injection] field '{fname}' is empty or not available, skipping targeted replacement.")
            continue
            
        serialized = _serialize_value(value, ftype)
        populated_fields.add(fname)
        token = field.get("template_token") or field.get("token") or f"{{{{{fname}}}}}"
        clean_token = token
        if token.upper().startswith("MERGEFIELD "):
            clean_token = token[len("MERGEFIELD "):].strip()

        # Handle table merge rows separately
        inj = field.get("injection_details") or {}
        injection_type = inj.get("injection_type", "text_placeholder")
        render_contract = field.get("render_contract") or {}
        render_strategy = render_contract.get("render_strategy", "")
        if injection_type == "text_placeholder" or not injection_type:
            if render_strategy == "mailmerge_table_region":
                injection_type = "table_merge_row"

        if ftype == "array" and field.get("source_block_ids"):
            if _inject_array_targeted(doc, field, value):
                targeted_injected_fields.add(fname)
                print(f"  [Targeted Array] successfully replaced {fname} using source block mappings")
            else:
                print(f"  [Targeted Array] no replacements performed for {fname}")
            continue

        if ftype in ("array", "array_object") and render_strategy == "repeat_block" and field.get("source_block_ids"):
            if _inject_repeat_block_targeted(doc, field, value):
                targeted_injected_fields.add(fname)
                print(f"  [Targeted RepeatBlock] successfully replaced {fname} using source block mappings")
            else:
                print(f"  [Targeted RepeatBlock] no replacements performed for {fname}")
            continue

        if injection_type == "table_merge_row":
            continue

        if ftype not in ("array", "array_object") and serialized:
            # Try original token, clean token, and stripped bracket token
            tokens_to_try = [clean_token, token]
            if clean_token.startswith("[") and clean_token.endswith("]"):
                tokens_to_try.append(clean_token[1:-1].strip())
            if token.startswith("[") and token.endswith("]"):
                tokens_to_try.append(token[1:-1].strip())

            for tok_variant in tokens_to_try:
                if _inject_targeted_value(doc, fname, tok_variant, serialized, field):
                    targeted_injected_fields.add(fname)
                    print(f"  [Targeted Injection] successfully replaced {fname} in specified blocks using '{tok_variant}'")
                    break

    # 2. Build mapping tables for global fallback pass
    for field in manifest_fields:
        fname = field.get("name", "")

        ftype = field.get("field_type", "scalar")
        inj = field.get("injection_details") or {}
        injection_type = inj.get("injection_type", "text_placeholder")
        render_contract = field.get("render_contract") or {}
        render_strategy = render_contract.get("render_strategy", "")
        if injection_type == "text_placeholder" or not injection_type:
            if render_strategy == "mailmerge_table_region":
                injection_type = "table_merge_row"

        # Skip global fallback for targeted fields, except table merge rows
        # which are intentionally built in this pass.
        if field.get("source_block_ids") and injection_type != "table_merge_row":
            print(f"  [Skip Global Fallback] field '{fname}' has source_block_ids, skipping global fallback to prevent collision.")
            continue

        value = extracted.get(fname)
        
        # User constraint: if value is empty or not available, do not inject/replace!
        if value is None or value == "" or value == [] or (isinstance(value, str) and not value.strip()):
            print(f"  [Skip Injection] field '{fname}' is empty or not available, skipping global fallback.")
            continue
            
        serialized = _serialize_value(value, ftype)
        populated_fields.add(fname)

        token = field.get("template_token") or field.get("token") or f"{{{{{fname}}}}}"
        clean_token = token
        if token.upper().startswith("MERGEFIELD "):
            clean_token = token[len("MERGEFIELD "):].strip()

        if injection_type == "table_merge_row":
            table_name = inj.get("table_name", "") or render_contract.get("region_name") or fname
            if table_name:
                raw_array = value or []
                if isinstance(raw_array, list):
                    items_mapped = []
                    for item in raw_array:
                        if isinstance(item, dict):
                            mapped_item = {}
                            sub_fields = field.get("sub_fields") or []
                            for skey, sval in item.items():
                                skey_norm = _normalize_field_name(skey)
                                sf_match = None
                                for sf in sub_fields:
                                    if sf.get("name") == skey_norm:
                                        sf_match = sf
                                        break
                                
                                xml_sf_name = None
                                if sf_match:
                                    sf_token = sf_match.get("template_token", "")
                                    if sf_token.upper().startswith("MERGEFIELD "):
                                        xml_sf_name = sf_token[len("MERGEFIELD "):].strip()
                                    else:
                                        xml_sf_name = sf_token
                                
                                if not xml_sf_name:
                                    xml_sf_name = skey
                                
                                mapped_item[xml_sf_name] = _serialize_value(sval, "scalar")
                            items_mapped.append(mapped_item)
                        elif item:
                            sub_fields = field.get("sub_fields") or []
                            xml_sf_name = "CheckType"
                            if sub_fields:
                                sf_token = sub_fields[0].get("template_token", "")
                                if sf_token.upper().startswith("MERGEFIELD "):
                                    xml_sf_name = sf_token[len("MERGEFIELD "):].strip()
                                else:
                                    xml_sf_name = sf_token
                            items_mapped.append({xml_sf_name: _serialize_value(item, "scalar")})
                    
                    table_data_map[table_name] = items_mapped
                    print(f"  [Dispatch] {fname} -> table_merge_row '{table_name}' ({len(items_mapped)} items) = {repr(serialized[:60])}")

        else:
            if ftype not in ("array", "array_object"):
                # Put in mergefield map (if it's not a visual bracket visual layout token)
                if not clean_token.startswith("[") and not clean_token.startswith("{"):
                    mergefield_map[clean_token] = serialized
                    mergefield_map[clean_token.lower()] = serialized
                    mergefield_map[clean_token.strip()] = serialized
                    mergefield_map[re.sub(r"[^a-zA-Z0-9]+", "", clean_token)] = serialized

                # Put in text placeholder map as fallback
                text_placeholder_map[token] = serialized
                text_placeholder_map[clean_token] = serialized
                text_placeholder_map[f"{{{{{fname}}}}}"] = serialized
                text_placeholder_map[f"{{{{ {fname} }}}}"] = serialized
                text_placeholder_map[f"[{fname}]"] = serialized
                text_placeholder_map[fname] = serialized
                
                # Unbracketed fallback for bracketed visual layout tokens
                if clean_token.startswith("[") and clean_token.endswith("]"):
                    unbracketed = clean_token[1:-1].strip()
                    text_placeholder_map[unbracketed] = serialized
                    text_placeholder_map[unbracketed.lower()] = serialized
                if token.startswith("[") and token.endswith("]"):
                    unbracketed = token[1:-1].strip()
                    text_placeholder_map[unbracketed] = serialized
                    text_placeholder_map[unbracketed.lower()] = serialized

                print(f"  [Dispatch] {fname} -> mergefield/text '{clean_token}' = {repr(serialized[:60])}")

    # Apply strategies
    print(f"[Injector] TableMergeRows to expand: {list(table_data_map.keys())}")
    print(f"[Injector] MergeFields to inject: {list(mergefield_map.keys())}")
    print(f"[Injector] Text placeholders to inject: {list(text_placeholder_map.keys())}")

    # Process table rows expansion first (since it duplicates mergefield markers)
    _inject_table_merge_rows(doc, table_data_map)
    _inject_array_paragraphs(doc, extracted, manifest_fields)
    _inject_mergefields(doc, mergefield_map)
    _inject_text_placeholders(doc, text_placeholder_map)

    # Final guardrail: clear any remaining known manifest tokens so unresolved
    # placeholders/mergefields do not leak into rendered output.
    unresolved_mergefields: dict[str, str] = {}
    unresolved_text_tokens: dict[str, str] = {}
    preserved_complex_tokens: list[str] = []
    for field in manifest_fields:
        fname = str(field.get("name") or "")
        if not fname or fname not in populated_fields:
            # Keep unresolved markers visible when the source value was missing.
            token = str(field.get("template_token") or field.get("token") or "").strip()
            if token:
                preserved_complex_tokens.append(token)
            for sub in field.get("sub_fields") or []:
                stoken = str(sub.get("template_token") or "").strip()
                if stoken:
                    preserved_complex_tokens.append(stoken)
            continue

        token = str(field.get("template_token") or field.get("token") or "").strip()
        if token:
            if token.upper().startswith("MERGEFIELD "):
                mf = token[len("MERGEFIELD "):].strip()
                if mf:
                    unresolved_mergefields[mf] = ""
            else:
                unresolved_text_tokens[token] = ""
                if token.startswith("[") and token.endswith("]"):
                    unresolved_text_tokens[token[1:-1].strip()] = ""
                # Bare token names (e.g., CandidateID) can still back a MERGEFIELD
                # display value, so include them in mergefield cleanup as well.
                if token and not token.startswith("[") and not token.startswith("{"):
                    unresolved_mergefields[token] = ""

        for sub in field.get("sub_fields") or []:
            stoken = str(sub.get("template_token") or "").strip()
            if not stoken:
                continue
            if stoken.upper().startswith("MERGEFIELD "):
                mf = stoken[len("MERGEFIELD "):].strip()
                if mf:
                    unresolved_mergefields[mf] = ""
            else:
                unresolved_text_tokens[stoken] = ""
                if stoken.startswith("[") and stoken.endswith("]"):
                    unresolved_text_tokens[stoken[1:-1].strip()] = ""
                if stoken and not stoken.startswith("[") and not stoken.startswith("{"):
                    unresolved_mergefields[stoken] = ""

    if unresolved_mergefields:
        _inject_mergefields(doc, unresolved_mergefields)
    if unresolved_text_tokens:
        _inject_text_placeholders(doc, unresolved_text_tokens)
    if preserved_complex_tokens:
        print(
            "[Injector] Preserved complex tokens for visibility (not auto-cleared): "
            f"{sorted(set(t for t in preserved_complex_tokens if t))}"
        )

    out_io = BytesIO()
    doc.save(out_io)
    return out_io.getvalue()


def inject_render_payload_into_docx(template_bytes: bytes, payload: dict, manifest: dict) -> bytes:
    fields = manifest.get("fields", []) if isinstance(manifest, dict) else []
    data = {}
    data.update(payload.get("render_values", {}))
    
    # Safely unpack placeholder_values list of dicts or fallback to dict
    ph_vals = payload.get("placeholder_values", {})
    if isinstance(ph_vals, list):
        for item in ph_vals:
            if isinstance(item, dict):
                if "name" in item:
                    data[item["name"]] = item.get("value")
                if "token" in item:
                    data[item["token"]] = item.get("value")
    elif isinstance(ph_vals, dict):
        data.update(ph_vals)

    for block_name, items in (payload.get("repeat_blocks", {}) or {}).items():
        data[block_name] = items
    return inject_data_into_docx(template_bytes, data, fields)
